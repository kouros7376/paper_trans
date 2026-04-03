# -*- coding: utf-8 -*-
"""
unified_converter.py
====================
다우오피스 전자결재 양식 통합 변환기

지원 파일 형식:
  - HWP  (.hwp)   : olefile + zlib 기반 바이너리 파싱
  - DOCX (.docx)  : python-docx 라이브러리
  - Excel (.xlsx, .xls) : openpyxl 라이브러리
  - PDF  (.pdf)   : pdfplumber 라이브러리

동작 순서:
  1. 입력 파일 확장자를 자동 감지
  2. 해당 형식에 맞는 파서로 텍스트/표 구조 추출
  3. 다우오피스 전자결재용 HTML 양식으로 변환
  4. HTML 파일 저장 (→ draft_approval.py에서 API 전송 가능)

사용법:
  python unified_converter.py "파일경로.hwp"
  python unified_converter.py "파일경로.docx"
  python unified_converter.py "파일경로.xlsx"
  python unified_converter.py "파일경로.pdf"

필요 패키지:
  pip install olefile python-docx openpyxl pdfplumber
"""

import sys
import struct
import zlib
import logging
from pathlib import Path

# ─────────────────────────────────────────────
# 로깅 설정
# ─────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger(__name__)

# 지원하는 파일 확장자 목록
SUPPORTED_EXTENSIONS = {'.hwp', '.docx', '.xlsx', '.xls', '.pdf'}


# ═════════════════════════════════════════════
# 공통 유틸리티
# ═════════════════════════════════════════════

def html_escape(text: str) -> str:
    """HTML 특수문자를 이스케이프 처리합니다."""
    return (str(text)
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;'))


import re as _re_common

def _is_cell_label(cell_text: str) -> bool:
    """
    셀 텍스트가 양식의 라벨(항목명)인지 판정합니다.
    모든 파서(HWP, DOCX, Excel)에서 공통으로 사용됩니다.

    라벨 예시: "소 속", "직 위", "일        자", "인계자", "서류",
              "인수∙인계 프로젝트", "인\n계\n자", "1. 인적사항"
    값 예시: "홍길동", "전략설계1본부", "010-1234-5678", "6,468,294원"
    """
    if not cell_text or not cell_text.strip():
        return False

    text = cell_text.strip()
    # 줄바꿈 제거한 텍스트 (세로 텍스트 정규화)
    flat = text.replace('\n', '').replace(' ', '')

    # ── 값으로 확정되는 패턴 (먼저 체크) ──
    # 숫자/영문/이메일/@/URL 포함 → 값 (단, 섹션 번호 "1." 제외)
    if _re_common.search(r'[a-zA-Z@]', flat):
        return False
    # 순수 숫자 또는 숫자+특수문자 (전화번호, 금액 등)
    if _re_common.match(r'^[\d,.\-+()/원년월일~\s]+$', flat):
        return False
    # 15글자 이상 긴 텍스트 → 값 (문장)
    if len(flat) > 15:
        return False

    # ── 라벨로 확정되는 패턴 ──
    # 1) 세로 텍스트: 한 글자씩 줄바꿈 (예: "인\n계\n자", "선\n결")
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if len(lines) >= 2 and all(len(l) <= 2 for l in lines):
        return True

    # 2) 한글 사이 공백 2칸 이상 (예: "일        자", "제  목", "기  타")
    if _re_common.search(r'[가-힣]  +[가-힣]', text):
        return True

    # 3) 한글 사이 공백 1칸 + 정규화 6글자 이하 (예: "소 속", "직 위", "휴가 종류")
    if _re_common.search(r'[가-힣] [가-힣]', text) and len(flat) <= 6:
        return True

    # 4) 섹션 번호 패턴 (예: "1. 인적사항", "2. 업무 인수∙인계 사항")
    if _re_common.match(r'^\d+\.\s', text):
        return True

    # 5) 짧은 한글 복합명사 (2~6글자, 숫자/영문 없음)
    #    예: "서류", "인계자", "담당자", "사고일시", "인수∙인계"
    if 2 <= len(flat) <= 6 and _re_common.match(r'^[가-힣∙·()/\s]+$', flat):
        return True

    return False


def _is_section_header(cell_text: str) -> bool:
    """섹션 제목인지 판정 (예: "1. 인적사항", "주 요 내 용")"""
    text = cell_text.strip()
    if _re_common.match(r'^\d+\.\s', text):
        return True
    flat = text.replace(' ', '')
    if 3 <= len(flat) <= 8 and _re_common.search(r'[가-힣]  +[가-힣]', text):
        return True
    return False


def build_daou_html(title: str, tables: list, paragraphs: list = None) -> str:
    """
    추출된 데이터를 다우오피스 전자결재용 HTML로 조립합니다.

    이 함수가 모든 파일 형식의 최종 HTML 출력을 담당합니다.
    다우오피스 필수 태그(data-id="appContent", apprPostParam 등)를 자동 삽입합니다.

    Args:
        title: 문서 제목
        tables: 테이블 데이터 리스트. 각 테이블은 딕셔너리:
                {
                    "caption": "섹션 제목 (선택)",
                    "rows": [
                        {"label": "라벨", "value": "값", "type": "input|textarea|text"}
                    ]
                }
        paragraphs: 테이블에 속하지 않는 일반 문단 리스트 (선택)

    Returns:
        다우오피스 전자결재용 HTML 문자열
    """
    # 파라미터 카운터 (다우오피스 필드 식별용)
    param_counter = 0

    html = []
    html.append('<!-- 다우오피스 전자결재 양식 - unified_converter.py 자동 생성 -->')
    html.append('<div data-id="appContent">')
    html.append('')
    html.append('<div style="font-family: \'Malgun Gothic\', dotum, Arial, sans-serif; '
                'font-size: 10pt; line-height: 1.6; margin: 0 auto; max-width: 800px;">')
    html.append('')

    # ── 문서 제목 ──
    html.append('  <div data-id="appTitle" style="text-align: center; margin-bottom: 20px;">')
    html.append('    <h2 style="font-size: 16pt; font-weight: bold; letter-spacing: 4px; '
                'border-bottom: 2px solid #333; padding-bottom: 8px;">')
    html.append(f'      {html_escape(title)}')
    html.append('    </h2>')
    html.append('  </div>')
    html.append('')

    # ── 테이블 섹션 ──
    for table_data in tables:
        caption = table_data.get("caption", "")
        rows = table_data.get("rows", [])
        # raw_rows: 원본 테이블의 행/열 구조를 그대로 보존한 데이터
        raw_rows = table_data.get("raw_rows", None)

        if not rows and not raw_rows:
            continue

        # ─── raw_rows 모드: 원본 셀 구조 그대로 HTML 변환 ───
        if raw_rows is not None:
            total_cols = table_data.get("total_cols", 1)
            html.append('  <table border="1" cellpadding="6" cellspacing="0"')
            html.append('         style="width: 100%; border-collapse: collapse; margin-bottom: 16px;">')
            if caption:
                html.append('    <thead>')
                html.append('      <tr style="background-color: #e6e8eb;">')
                html.append(f'        <th colspan="{total_cols}" style="text-align: left; padding: 6px 10px; font-size: 10pt;">')
                html.append(f'          {html_escape(caption)}')
                html.append('        </th>')
                html.append('      </tr>')
                html.append('    </thead>')
            html.append('    <tbody>')

            for row_cells in raw_rows:
                html.append('      <tr>')
                for cell in row_cells:
                    text = cell.get("text", "")
                    colspan = cell.get("colspan", 1)
                    rowspan = cell.get("rowspan", 1)
                    is_label = cell.get("is_label", False)
                    escaped = html_escape(text)

                    # colspan/rowspan 속성 문자열
                    span_attr = ""
                    if colspan > 1:
                        span_attr += f' colspan="{colspan}"'
                    if rowspan > 1:
                        span_attr += f' rowspan="{rowspan}"'

                    if is_label:
                        # 라벨 셀 (회색 배경, 굵은 글씨)
                        html.append(f'        <td{span_attr} style="background-color: #f5f6f8; font-weight: bold; '
                                  f'text-align: center; padding: 6px 10px; white-space: pre-line;">')
                        html.append(f'          {escaped}')
                        html.append('        </td>')
                    elif text.strip():
                        # 값 셀 (편집 가능한 입력 필드)
                        param_counter += 1
                        param_id = f'apprPostParam{param_counter}'
                        # 여러 줄이면 textarea, 한 줄이면 input
                        if '\n' in text or len(text) > 100:
                            html.append(f'        <td{span_attr} style="padding: 6px 10px; vertical-align: top;">')
                            html.append(f'          <textarea name="field_{param_counter}" data-id="{param_id}" '
                                      f'style="width: 98%; min-height: 60px; border: none; '
                                      f'font-family: inherit; font-size: 10pt; resize: vertical;"'
                                      f'>{escaped}</textarea>')
                            html.append('        </td>')
                        else:
                            html.append(f'        <td{span_attr} style="padding: 6px 10px;">')
                            html.append(f'          <input type="text" name="field_{param_counter}" '
                                      f'data-id="{param_id}" value="{escaped}" '
                                      f'style="width: 95%; border: none; font-family: inherit; font-size: 10pt;" />')
                            html.append('        </td>')
                    else:
                        # 빈 셀
                        html.append(f'        <td{span_attr} style="padding: 6px 10px;">&nbsp;</td>')

                html.append('      </tr>')
            html.append('    </tbody>')
            html.append('  </table>')
            html.append('')
            continue

        # ─── 기존 rows 모드 (호환) ───
        max_cols = 2
        for r in rows:
            if r.get("type") == "multi_cell":
                max_cols = max(max_cols, len(r.get("cells", [])))

        html.append('  <table border="1" cellpadding="6" cellspacing="0"')
        html.append('         style="width: 100%; border-collapse: collapse; margin-bottom: 16px;">')
        if caption:
            html.append('    <thead>')
            html.append('      <tr style="background-color: #e6e8eb;">')
            html.append(f'        <th colspan="{max_cols}" style="text-align: left; padding: 6px 10px; font-size: 10pt;">')
            html.append(f'          {html_escape(caption)}')
            html.append('        </th>')
            html.append('      </tr>')
            html.append('    </thead>')

        html.append('    <tbody>')

        for row in rows:
            label = row.get("label")
            value = row.get("value", "")
            row_type = row.get("type", "input")

            param_counter += 1
            param_id = f'apprPostParam{param_counter}'
            escaped_value = html_escape(value)

            if row_type == "header":
                html.append('      <tr style="background-color: #f0f1f3;">')
                html.append(f'        <td colspan="{max_cols}" style="font-weight: bold; padding: 6px 10px;">')
                html.append(f'          {escaped_value}')
                html.append('        </td>')
                html.append('      </tr>')

            elif row_type == "multi_cell":
                cells = row.get("cells", [])
                is_header = row.get("is_header", False)
                html.append('      <tr>')
                for ci, cell_val in enumerate(cells):
                    escaped_cv = html_escape(cell_val)
                    if is_header or (ci % 2 == 0 and ci + 1 < len(cells)):
                        html.append(f'        <td style="background-color: #f5f6f8; font-weight: bold; '
                                  f'text-align: center; padding: 6px 10px;">')
                        html.append(f'          {escaped_cv}')
                        html.append('        </td>')
                    else:
                        param_counter += 1
                        param_id = f'apprPostParam{param_counter}'
                        html.append(f'        <td style="padding: 6px 10px;">')
                        html.append(f'          <input type="text" name="field_{param_counter}" '
                                  f'data-id="{param_id}" value="{escaped_cv}" '
                                  f'style="width: 95%; border: none; font-family: inherit; font-size: 10pt;" />')
                        html.append('        </td>')
                html.append('      </tr>')

            elif row_type == "text" or label is None:
                html.append('      <tr>')
                html.append(f'        <td colspan="{max_cols}" style="text-align: center; padding: 10px;">')
                html.append(f'          {escaped_value}')
                html.append('        </td>')
                html.append('      </tr>')

            elif row_type == "textarea":
                val_span = max_cols - 1
                html.append('      <tr>')
                html.append(f'        <td style="background-color: #f5f6f8; font-weight: bold; '
                          f'text-align: center; width: 25%; vertical-align: top;">')
                html.append(f'          {html_escape(label)}')
                html.append('        </td>')
                html.append(f'        <td colspan="{val_span}">')
                html.append(f'          <textarea name="field_{param_counter}" data-id="{param_id}" '
                          f'style="width: 98%; min-height: 80px; border: none; '
                          f'font-family: inherit; font-size: 10pt; resize: vertical;"'
                          f'>{escaped_value}</textarea>')
                html.append('        </td>')
                html.append('      </tr>')

            else:
                val_span = max_cols - 1
                html.append('      <tr>')
                html.append(f'        <td style="background-color: #f5f6f8; font-weight: bold; '
                          f'text-align: center; width: 25%;">')
                html.append(f'          {html_escape(label)}')
                html.append('        </td>')
                html.append(f'        <td colspan="{val_span}">')
                html.append(f'          <input type="text" name="field_{param_counter}" '
                          f'data-id="{param_id}" value="{escaped_value}" '
                          f'style="width: 95%; border: none; font-family: inherit; font-size: 10pt;" />')
                html.append('        </td>')
                html.append('      </tr>')

        html.append('    </tbody>')
        html.append('  </table>')
        html.append('')

    # ── 일반 문단 (테이블 외 텍스트) → 제출문구/날짜/서명란 자동 감지 ──
    if paragraphs:
        import re as _re_para

        # 제출 문구 패턴: "위와 같이 ~합니다" / "상기와 같이 ~합니다"
        submit_pattern = _re_para.compile(r'(위와\s*같이|상기와\s*같이).+(합니다|제출합니다|신청합니다)')
        # 날짜 패턴: "20XX년 XX월 XX일" 또는 "20년  월  일"
        date_pattern = _re_para.compile(r'\d{2,4}\s*년\s+\d{0,2}\s*월\s+\d{0,2}\s*일')
        # 서명란 패턴: "OOO : OOO (인)" 또는 "OOO :      (인)"
        sign_pattern = _re_para.compile(r'(.+?)\s*:\s*(.*?)\s*\(인\)')

        # 제출문구/날짜/서명란을 하나의 테이블로 묶어서 출력
        submit_rows = []   # 제출 관련 문단
        other_rows = []    # 그 외 일반 문단

        for para in paragraphs:
            stripped = para.strip()
            if not stripped:
                continue

            # 이미지 마커는 정규식 검사 없이 바로 일반 문단으로 분류
            if stripped.startswith('__IMG__'):
                other_rows.append(stripped)
                continue

            if submit_pattern.search(stripped):
                submit_rows.append({"type": "submit_text", "value": stripped})
            elif date_pattern.search(stripped):
                submit_rows.append({"type": "submit_date", "value": stripped})
            elif sign_pattern.search(stripped):
                # 여러 서명란이 줄바꿈으로 합쳐진 경우 분리
                for line in stripped.split('\n'):
                    m = sign_pattern.search(line.strip())
                    if m:
                        submit_rows.append({
                            "type": "sign",
                            "role": m.group(1).strip(),
                            "name": m.group(2).strip(),
                        })
                    elif line.strip():
                        submit_rows.append({"type": "submit_text", "value": line.strip()})
            else:
                other_rows.append(stripped)

        # 일반 문단 먼저 출력
        for text in other_rows:
            # 이미지 마커 처리
            if text.startswith('__IMG__'):
                img_src = text[7:]
                html.append(f'  <div style="text-align: center; margin: 16px 0;">')
                html.append(f'    <img src="{img_src}" style="max-width: 200px; height: auto;" />')
                html.append(f'  </div>')
                continue
            param_counter += 1
            html.append(f'  <div style="margin-bottom: 8px;">')
            html.append(f'    <p>{html_escape(text)}</p>')
            html.append(f'  </div>')

        # 제출문구/서명란 테이블 출력
        if submit_rows:
            html.append('')
            html.append('  <!-- ===== 제출 문구 및 서명란 ===== -->')
            # 제출 문구 (중앙 정렬 텍스트)
            for row in submit_rows:
                if row["type"] == "submit_text":
                    html.append(f'  <div style="text-align: center; margin-bottom: 10px; font-size: 10pt;">')
                    html.append(f'    {html_escape(row["value"])}')
                    html.append(f'  </div>')
                elif row["type"] == "submit_date":
                    param_counter += 1
                    param_id = f'apprPostParam{param_counter}'
                    html.append(f'  <div style="text-align: center; margin-bottom: 20px; font-size: 10pt;">')
                    html.append(f'    <input type="text" name="submitDate" data-id="{param_id}" '
                              f'value="{html_escape(row["value"])}" '
                              f'style="border: none; text-align: center; font-family: inherit; font-size: 10pt;" />')
                    html.append(f'  </div>')

            # 서명란 테이블
            sign_rows = [r for r in submit_rows if r["type"] == "sign"]
            if sign_rows:
                html.append('')
                html.append('  <table border="0" cellpadding="6" cellspacing="0"')
                html.append('         style="width: 60%; margin: 0 auto; font-size: 10pt;">')
                html.append('    <tbody>')
                for srow in sign_rows:
                    param_counter += 1
                    param_id = f'apprPostParam{param_counter}'
                    escaped_name = html_escape(srow["name"])
                    escaped_role = html_escape(srow["role"])
                    html.append('      <tr>')
                    html.append(f'        <td style="text-align: right; width: 40%;">{escaped_role} :</td>')
                    html.append(f'        <td style="border-bottom: 1px solid #333; width: 45%;">')
                    html.append(f'          <input type="text" name="sig_{escaped_role}" data-id="{param_id}" '
                              f'value="{escaped_name}" '
                              f'style="border: none; width: 90%; font-family: inherit; font-size: 10pt;" />')
                    html.append(f'        </td>')
                    html.append(f'        <td style="width: 15%; text-align: center;">(인)</td>')
                    html.append('      </tr>')
                html.append('    </tbody>')
                html.append('  </table>')
            html.append('')

    html.append('</div><!-- /font-family div -->')
    html.append('</div><!-- /data-id="appContent" -->')

    return '\n'.join(html)


# ═════════════════════════════════════════════
# [1] HWP 파서
# ═════════════════════════════════════════════

def _parse_hwp(file_path: Path) -> tuple:
    """
    HWP 파일에서 테이블 셀 구조를 포함하여 텍스트를 추출합니다.

    핵심 원리:
      - HWP 레코드의 Level 필드로 계층 구조를 파악
      - LIST_HEADER(72) 태그가 나올 때마다 새로운 셀 시작
      - 셀 내부의 PARA_TEXT(67) 태그에서 텍스트 추출
      - 결재란/워터마크 등 불필요 셀은 자동 필터링

    Returns:
        (title, tables, paragraphs) 튜플
    """
    import olefile

    ole = olefile.OleFileIO(str(file_path))

    # 압축/암호화 여부 확인
    header = ole.openstream('FileHeader').read()
    flags = int.from_bytes(header[36:40], 'little')
    is_compressed = bool(flags & 0x01)
    if flags & 0x02:
        ole.close()
        raise ValueError("암호화된 HWP 파일은 지원하지 않습니다.")

    # ── BinData에서 이미지 추출 (로고 등) ──
    import base64 as _b64
    embedded_images = []  # base64 인코딩된 이미지 목록
    for stream in ole.listdir():
        path = '/'.join(stream)
        if path.startswith('BinData/') and any(path.lower().endswith(ext) for ext in ('.jpg', '.jpeg', '.png', '.gif', '.bmp')):
            try:
                img_data = ole.openstream(path).read()
                # 압축된 이미지 → zlib 해제 시도
                try:
                    img_data = zlib.decompress(img_data, -15)
                except zlib.error:
                    pass  # 이미 비압축 상태
                ext = path.rsplit('.', 1)[-1].lower()
                mime = {'jpg': 'jpeg', 'jpeg': 'jpeg', 'png': 'png', 'gif': 'gif', 'bmp': 'bmp'}.get(ext, 'jpeg')
                b64 = _b64.b64encode(img_data).decode('ascii')
                embedded_images.append(f'data:image/{mime};base64,{b64}')
            except Exception:
                pass

    # BodyText 섹션 추출 → 셀 단위로 파싱
    all_tables = []  # 표 목록 (각 표 = 셀 텍스트 리스트)
    body_paragraphs = []  # 표 밖의 본문 문단

    for stream in ole.listdir():
        path = '/'.join(stream)
        if path.startswith('BodyText/Section'):
            raw = ole.openstream(path).read()
            data = zlib.decompress(raw, -15) if is_compressed else raw
            tables_in_section, paras_in_section = _hwp_parse_with_structure(data)
            all_tables.extend(tables_in_section)
            body_paragraphs.extend(paras_in_section)
    ole.close()

    # 워터마크/불필요 텍스트 필터
    # 워터마크 텍스트 (정확 일치만 사용하여 오탐 방지)
    SKIP = {'문서서식포탈비', '문서서식포탈비즈폼', '비즈폼'}

    # ── 표 데이터를 다우오피스 HTML 구조로 변환 ──
    title = file_path.stem
    title_found = False  # 제목이 이미 추출되었는지 여부
    result_tables = []

    # 결재란/워터마크에서 자주 나오는 키워드
    APPROVAL_KEYWORDS = {'결', '재', '결재', '결\n재', '담   당', '이   사', '본 부 장',
                        '담당', '이사', '본부장', '사장', '대표이사', '부사장', '전무', '상무'}

    for table_rows in all_tables:
        # table_rows = [[{text, colspan, rowspan}, ...], ...] (새 형식: 딕셔너리)

        # 전체 셀 텍스트 추출 (워터마크/결재란 검사용)
        all_cell_texts = []
        for row in table_rows:
            for cell in row:
                t = cell["text"].strip() if isinstance(cell, dict) else str(cell).strip()
                if t:
                    all_cell_texts.append(t)

        clean_texts = [t for t in all_cell_texts if not any(s in t for s in SKIP)]
        if not clean_texts:
            continue

        # 결재란 행 제거
        filtered_rows = []
        for row_cells in table_rows:
            row_texts = []
            for cell in row_cells:
                t = cell["text"].strip() if isinstance(cell, dict) else str(cell).strip()
                if t and not any(s in t for s in SKIP):
                    row_texts.append(t)
            approval_in_row = sum(1 for c in row_texts if c in APPROVAL_KEYWORDS)
            if approval_in_row >= 2:
                # 결재란 행 → 건너뛰되 제목 추출 시도
                for c in row_texts:
                    if c not in APPROVAL_KEYWORDS and len(c) < 30:
                        candidate = c.replace('\n', '').replace(' ', '')
                        if 2 <= len(candidate) <= 15:
                            title = c
                            title_found = True
                            break
                continue
            filtered_rows.append(row_cells)

        if not filtered_rows:
            continue

        # ── HWP 바이너리에서 읽은 정확한 colspan/rowspan을 그대로 사용 ──
        total_cols = 0
        for row in filtered_rows:
            row_cols = sum(cell.get("colspan", 1) if isinstance(cell, dict) else 1 for cell in row)
            total_cols = max(total_cols, row_cols)
        if total_cols == 0:
            total_cols = 1

        raw_rows = []
        for row_cells in filtered_rows:
            html_row = []
            has_content = False
            for cell in row_cells:
                if isinstance(cell, dict):
                    text = cell["text"].strip()
                    # 워터마크 텍스트 제거
                    if any(s in text for s in SKIP):
                        text = ""
                    cs = cell.get("colspan", 1)
                    rs = cell.get("rowspan", 1)
                else:
                    text = str(cell).strip()
                    cs = 1
                    rs = 1

                if text:
                    has_content = True

                html_row.append({
                    "text": text,
                    "colspan": cs,
                    "rowspan": rs,
                    "is_label": _is_cell_label(text),
                })

            # 제목 감지
            non_empty = [c for c in html_row if c["text"]]
            if not title_found and len(non_empty) == 1:
                t = non_empty[0]["text"]
                flat = t.replace('\n', '').replace(' ', '')
                if 2 <= len(flat) <= 15 and '  ' in t:
                    title = t
                    title_found = True
                    continue

            if has_content:
                raw_rows.append(html_row)

        if raw_rows:
            result_tables.append({
                "caption": "",
                "rows": [],
                "raw_rows": raw_rows,
                "total_cols": total_cols,
            })

    # 이미지를 paragraphs 뒤에 추가 (build_daou_html에서 처리)
    if embedded_images:
        for img_src in embedded_images:
            body_paragraphs.append(f'__IMG__{img_src}')

    return title, result_tables, body_paragraphs


def _hwp_parse_with_structure(section_data: bytes) -> tuple:
    """
    HWP 바이너리 레코드를 Level + LIST_HEADER 기반으로 파싱하여
    표의 행/열 구조를 정확히 추출합니다.

    LIST_HEADER(72) 페이로드 구조 (오프셋, uint16 LE):
      - [0]: 문단 수
      - [4]: col (열 위치, 0부터)
      - [5]: row (행 위치, 0부터)
      - [6]: col_span (열 병합 수)
      - [7]: row_span (행 병합 수)

    Returns:
        (tables, paragraphs) 튜플
        - tables: 각 표는 {(row, col): text, ...} 딕셔너리 → 행 단위 리스트로 변환
        - paragraphs: 표 밖 문단 목록
    """
    EXTENDED = {1, 2, 3, 4, 5, 6, 7, 8, 9, 11, 15, 16, 17, 18, 21, 22, 23, 24}
    pos = 0

    tables = []           # 결과: 표 목록 (각 표 = 행 리스트의 리스트)
    paragraphs = []       # 결과: 표 밖 문단

    # 현재 파싱 상태
    current_table = {}    # {(row, col): {"text": str, "colspan": int, "rowspan": int}}
    current_cell_text = []
    current_cell_pos = None  # (row, col)
    current_cell_span = (1, 1)  # (col_span, row_span)
    in_table = False
    max_row = 0
    max_col = 0

    def _flush_cell():
        """현재 셀의 텍스트와 병합 정보를 테이블 딕셔너리에 저장"""
        nonlocal current_cell_text, current_cell_pos
        if current_cell_pos is not None:
            text = '\n'.join(current_cell_text) if current_cell_text else ''
            current_table[current_cell_pos] = {
                "text": text,
                "colspan": current_cell_span[0],
                "rowspan": current_cell_span[1],
            }
        current_cell_text = []

    def _flush_table():
        """현재 테이블을 raw_rows 형식으로 변환하여 저장"""
        nonlocal current_table, in_table, max_row, max_col, current_cell_pos
        _flush_cell()
        if current_table:
            # 병합된 셀이 차지하는 위치를 추적 (중복 출력 방지)
            merged_positions = set()
            for (r, c), info in current_table.items():
                cs = info.get("colspan", 1)
                rs = info.get("rowspan", 1)
                for dr in range(rs):
                    for dc in range(cs):
                        if dr != 0 or dc != 0:
                            merged_positions.add((r + dr, c + dc))

            # 행 단위 리스트로 변환 (병합 정보 포함)
            row_list = []
            for r in range(max_row + 1):
                row_cells = []
                for c in range(max_col + 1):
                    if (r, c) in merged_positions:
                        continue  # 병합으로 가려진 셀은 건너뛰기
                    info = current_table.get((r, c))
                    if info:
                        row_cells.append(info)
                    else:
                        row_cells.append({"text": "", "colspan": 1, "rowspan": 1})
                if row_cells:
                    row_list.append(row_cells)
            if row_list:
                tables.append(row_list)
        current_table = {}
        current_cell_pos = None
        in_table = False
        max_row = 0
        max_col = 0

    while pos + 4 <= len(section_data):
        hdr = struct.unpack_from('<I', section_data, pos)[0]
        tag = hdr & 0x3FF
        level = (hdr >> 10) & 0x3FF
        size = (hdr >> 20) & 0xFFF
        pos += 4
        if size == 0xFFF:
            if pos + 4 > len(section_data):
                break
            size = struct.unpack_from('<I', section_data, pos)[0]
            pos += 4

        # ── L0 PARA_HEADER: 표 밖의 새 문단 → 이전 표 완료 ──
        if tag == 66 and level == 0:
            if in_table:
                _flush_table()

        # ── L1 CTRL_HEADER(71): 표/그리기 시작 ──
        elif tag == 71 and level == 1:
            if not in_table:
                in_table = True
                current_table = {}
                current_cell_text = []
                current_cell_pos = None
                max_row = 0
                max_col = 0

        # ── L2 LIST_HEADER(72): 새 셀 시작 + 위치/병합 정보 파싱 ──
        elif tag == 72 and level >= 2:
            _flush_cell()
            # LIST_HEADER 페이로드에서 행/열 위치 + colspan/rowspan 추출
            payload = section_data[pos:pos + size]
            if size >= 16:  # 최소 8개 uint16 필요
                col = struct.unpack_from('<H', payload, 8)[0]       # offset 8: col
                row = struct.unpack_from('<H', payload, 10)[0]      # offset 10: row
                col_span = struct.unpack_from('<H', payload, 12)[0] # offset 12: col_span
                row_span = struct.unpack_from('<H', payload, 14)[0] # offset 14: row_span
                current_cell_pos = (row, col)
                current_cell_span = (max(col_span, 1), max(row_span, 1))
                # 병합 범위를 고려한 최대 행/열 계산
                if row + row_span - 1 > max_row:
                    max_row = row + row_span - 1
                if col + col_span - 1 > max_col:
                    max_col = col + col_span - 1
            else:
                current_cell_pos = None
                current_cell_span = (1, 1)

        # ── PARA_TEXT(67): 텍스트 추출 ──
        elif tag == 67:
            text = _hwp_decode_text(section_data[pos:pos + size], EXTENDED)
            if text.strip():
                if in_table and level >= 3:
                    current_cell_text.append(text.strip())
                elif level <= 1:
                    paragraphs.append(text.strip())

        pos += size

    # 마지막 표 처리
    if in_table:
        _flush_table()

    return tables, paragraphs


def _hwp_decode_text(text_data: bytes, extended_chars: set) -> str:
    """HWP PARA_TEXT 바이너리를 문자열로 디코딩합니다."""
    text = ''
    i = 0
    while i < len(text_data) - 1:
        cc = struct.unpack_from('<H', text_data, i)[0]
        i += 2
        if cc < 32:
            if cc in extended_chars:
                if cc == 9:
                    text += '\t'
                i += 14
            elif cc == 10:
                text += '\n'
        else:
            text += chr(cc)
    return text


def _match_hwp_cell_pairs(cells: list) -> list:
    """
    HWP 표의 셀 목록을 라벨-값 쌍으로 매칭합니다.

    HWP PrvText의 원본 구조를 참고하여:
    - 연속된 셀이 <라벨><값> 쌍인지 판별
    - 단독 셀(전체 너비)은 텍스트로 처리
    """
    import re
    rows = []
    i = 0

    while i < len(cells):
        text = cells[i]

        # 다음 셀이 있는 경우 → 라벨-값 쌍 판별
        if i + 1 < len(cells):
            next_text = cells[i + 1]

            # 현재 셀이 라벨처럼 보이는지 판단
            # 라벨 특성: 콜론으로 끝남, 한글+공백 패턴, 짧은 한글 복합어
            normalized = text.replace(' ', '')
            is_label = False

            # 콜론으로 끝나는 라벨
            if text.strip().endswith(':') and len(text.strip()) < 25:
                is_label = True
            # 한글만 + 짧은 텍스트 (숫자/영문 없음)
            elif (re.match(r'^[가-힣\s/·∙()]+$', text) and
                  len(normalized) <= 15 and
                  not re.search(r'[0-9a-zA-Z]', text)):
                # "위와 같이..." 같은 문장은 라벨이 아님
                if not any(kw in text for kw in ['같이', '합니다', '바랍니다', '제출']):
                    is_label = True
            # "항목 및 항목" 패턴
            elif re.match(r'^[가-힣\s]+및[가-힣\s]+$', text.strip()) and len(normalized) <= 15:
                is_label = True

            if is_label:
                row_type = "textarea" if '\n' in next_text or len(next_text) > 80 else "input"
                rows.append({"label": text, "value": next_text, "type": row_type})
                i += 2
                continue

        # 단독 셀 (전체 너비 텍스트)
        rows.append({"label": None, "value": text, "type": "text"})
        i += 1

    return rows


def _detect_label_value_pairs(texts: list) -> list:
    """
    텍스트 목록에서 라벨-값 쌍을 자동 감지합니다.

    양식 문서의 라벨 특성:
      - 한글 사이에 공백이 2개 이상 (예: "부    서", "성    명")
      - 복합 명사형 (예: "사고일시", "사고장소", "사진첨부")
      - 콜론(:)으로 끝나는 텍스트 (예: "작 성 자 :")
      - "~및~" 패턴 (예: "사고원인 및 사고내용")

    값과 구분하는 핵심 기준:
      - 라벨은 보통 "개념어"이고, 값은 "실제 데이터"
      - 라벨에는 숫자, 날짜, 전화번호 등이 포함되지 않음
      - "홍 길 동", "대 리" 같은 것은 값(사람이름/직위)이지 라벨이 아님

    이 함수는 범용적으로 동작하여 HWP/DOCX/PDF 모두에서 사용 가능합니다.
    """
    import re

    def _is_label(text: str) -> bool:
        """텍스트가 양식 라벨인지 판정합니다."""
        stripped = text.strip()
        normalized = stripped.replace(' ', '')

        # 숫자, 영문, 특수기호가 포함되면 라벨이 아닐 가능성 높음
        # (단, 콜론(:), 슬래시(/), 및, ·, ∙는 라벨에 포함 가능)
        if re.search(r'[0-9a-zA-Z]', normalized):
            return False

        # 콜론으로 끝나는 텍스트 → 라벨 확정
        if stripped.endswith(':') and len(normalized) < 20:
            return True

        # "~및~" 패턴 (예: "사고원인 및 사고내용")
        if re.match(r'^[가-힣\s]+및[가-힣\s]+$', stripped) and len(normalized) <= 15:
            return True

        # 한글만으로 구성된 경우 → 추가 판정 필요
        if not re.match(r'^[가-힣/·∙:]+$', normalized):
            return False

        # 글자 수가 너무 적으면 (1글자) → 라벨 아님 (결, 재 등 제외용)
        if len(normalized) <= 1:
            return False

        # 핵심 판정: 공백 패턴으로 라벨 vs 값 구분
        # 라벨 = 한글 사이에 공백 2개 이상 (예: "부    서", "성    명")
        # 값 = 한글 사이에 공백 1개 (예: "홍 길 동", "대 리")
        has_multi_space = bool(re.search(r'[가-힣]  +[가-힣]', stripped))

        # 공백이 없는 복합어 (예: "사고일시", "사고장소", "사진첨부")
        is_compound = len(normalized) >= 2 and ' ' not in stripped

        # 공백이 있지만 1칸만 (예: "홍 길 동") → 값으로 판정
        has_single_space_only = (
            ' ' in stripped and
            not has_multi_space and
            not stripped.endswith(':')
        )

        if has_single_space_only:
            return False

        # 라벨 후보: 다중 공백 패턴 또는 공백 없는 복합어
        if has_multi_space or is_compound:
            return True

        return False

    rows = []
    i = 0

    while i < len(texts):
        text = texts[i]

        # "사고원인 및" + "사고내용" 패턴 (라벨이 여러 줄에 걸치는 경우)
        normalized = text.replace(' ', '')
        if i + 1 < len(texts):
            next_norm = texts[i + 1].replace(' ', '')
            combined_norm = normalized + next_norm
            # 두 줄을 합치면 라벨이 되는 경우
            if (normalized in ('사고원인및', '사고원인') and next_norm == '사고내용'):
                label = '사고원인 및 사고내용'
                value = texts[i + 2] if i + 2 < len(texts) else ''
                row_type = "textarea" if '\n' in value or len(value) > 40 else "input"
                rows.append({"label": label, "value": value, "type": row_type})
                i += 3
                continue

        if _is_label(text):
            if i + 1 < len(texts):
                next_text = texts[i + 1]
                if _is_label(next_text):
                    # 다음 텍스트도 라벨 → 현재 라벨은 값 없음
                    rows.append({"label": text, "value": "", "type": "input"})
                    i += 1
                else:
                    # 라벨 + 값 쌍
                    row_type = "textarea" if '\n' in next_text or len(next_text) > 80 else "input"
                    rows.append({"label": text, "value": next_text, "type": row_type})
                    i += 2
            else:
                rows.append({"label": text, "value": "", "type": "input"})
                i += 1
        else:
            # 단독 텍스트 (값, 제출 문구 등)
            rows.append({"label": None, "value": text, "type": "text"})
            i += 1

    return rows


# ═════════════════════════════════════════════
# [2] DOCX 파서
# ═════════════════════════════════════════════

def _parse_docx(file_path: Path) -> tuple:
    """
    DOCX 파일에서 텍스트와 표를 추출합니다.

    mammoth 라이브러리로 DOCX → HTML 변환 후,
    다우오피스 양식에 맞게 래핑합니다.
    mammoth이 colspan/rowspan/이미지를 정확하게 처리합니다.

    Returns:
        (title, tables, paragraphs) 튜플
    """
    from docx import Document

    doc = Document(str(file_path))

    # ── 제목 감지 ──
    # 1순위: Heading 스타일 또는 가운데 정렬 문단
    # 2순위: 첫 번째 표의 첫 번째 단일 셀 (양식에서 제목이 표 안에 있는 경우)
    # 3순위: 첫 번째 비어있지 않은 문단
    title = ""
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') or (
            para.alignment is not None and para.alignment == 1  # CENTER
        ):
            if para.text.strip():
                title = para.text.strip()
                break
    # Heading이 없으면 첫 번째 비어있지 않은 문단 사용
    if not title:
        for para in doc.paragraphs:
            if para.text.strip():
                title = para.text.strip()
                break
    # 그래도 없으면 파일명 사용 (단, UUID가 아닌 경우만)
    if not title:
        import re as _re
        stem = file_path.stem
        # UUID 패턴(8자리 hex)이 아닌 경우만 파일명 사용
        if not _re.match(r'^[0-9a-f]{8}$', stem):
            title = stem
        else:
            title = "문서"

    tables = []
    extra_paragraphs = []

    # ── 표(Table) 추출 → raw_rows 방식으로 원본 구조 보존 ──
    for table_idx, table in enumerate(doc.tables):
        total_cols = len(table.columns) if table.columns else 1
        raw_rows = []

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]

            # 셀 병합 감지: python-docx는 병합된 셀을 동일 텍스트로 반복함
            html_row = []
            i = 0
            while i < len(cells):
                cell_text = cells[i]
                colspan = 1
                # 동일한 텍스트가 연속되면 colspan으로 처리
                while i + colspan < len(cells) and cells[i + colspan] == cell_text:
                    colspan += 1

                # 공통 라벨 판정 함수 사용
                is_label = _is_cell_label(cell_text)

                html_row.append({
                    "text": cell_text,
                    "colspan": colspan,
                    "rowspan": 1,
                    "is_label": is_label,
                })
                i += colspan

            # 완전히 빈 행은 건너뛰기
            if any(c["text"] for c in html_row):
                raw_rows.append(html_row)

        if raw_rows:
            tables.append({
                "caption": "",
                "rows": [],
                "raw_rows": raw_rows,
                "total_cols": total_cols,
            })

    # ── 본문 문단 추출 및 섹션 제목 → 테이블 캡션 연결 ──
    # python-docx의 doc.element.body를 순회하여 문단-테이블 순서를 파악
    import re as _re_sect
    table_idx = 0
    section_pattern = _re_sect.compile(r'^\d+\.\s+')  # "1. 신청자 정보" 패턴

    # 테이블에서 이미 추출된 텍스트를 수집 (중복 출력 방지)
    table_texts = set()
    for tbl in tables:
        # raw_rows 방식
        for rr in tbl.get("raw_rows", []):
            for cell in rr:
                if cell.get("text", "").strip():
                    table_texts.add(cell["text"].strip())
        # 기존 rows 방식 (호환)
        for r in tbl.get("rows", []):
            if r.get("value"):
                table_texts.add(r["value"].strip())
            if r.get("label"):
                table_texts.add(r["label"].strip())
            for cell_text in r.get("cells", []):
                if cell_text.strip():
                    table_texts.add(cell_text.strip())

    # 문서 본문 요소를 순서대로 순회하여 섹션 제목을 다음 테이블의 캡션으로 연결
    from docx.oxml.ns import qn as _qn
    pending_caption = ""
    tbl_counter = 0

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            # 문단(paragraph) 처리
            # w:r/w:t 태그에서 텍스트를 추출 (XML 네임스페이스 포함)
            ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            parts = []
            for r_elem in element.findall(f'{{{ns}}}r'):
                for t_elem in r_elem.findall(f'{{{ns}}}t'):
                    if t_elem.text:
                        parts.append(t_elem.text)
            full_text = ''.join(parts).strip()
            if not full_text:
                continue
            if full_text == title:
                continue

            # 섹션 제목 패턴("1. xxx") → 다음 테이블의 캡션으로 대기
            if section_pattern.match(full_text):
                pending_caption = full_text
            elif full_text not in table_texts:
                extra_paragraphs.append(full_text)

        elif tag == 'tbl':
            # 테이블 요소 → 대기 중인 캡션 연결
            if tbl_counter < len(tables) and pending_caption:
                tables[tbl_counter]["caption"] = pending_caption
                pending_caption = ""
            tbl_counter += 1

    # 표가 없으면 문단을 라벨-값 감지 시도
    if not tables and extra_paragraphs:
        rows = _detect_label_value_pairs(extra_paragraphs)
        tables = [{"caption": "", "rows": rows}]
        extra_paragraphs = []

    return title, tables, extra_paragraphs


# ═════════════════════════════════════════════
# [3] Excel 파서
# ═════════════════════════════════════════════

def _parse_excel(file_path: Path) -> tuple:
    """
    Excel 파일(.xlsx/.xls)에서 시트 데이터를 추출합니다.

    원본 레이아웃을 최대한 보존하기 위해:
      - 다열 구조를 HTML 테이블 행으로 그대로 변환
      - 병합 셀 처리 (중복 값 한 번만 표시)
      - 날짜 값 자동 포맷팅
      - 빈 행/열 자동 건너뛰기

    Returns:
        (title, tables, paragraphs) 튜플
    """
    from openpyxl import load_workbook
    from datetime import datetime

    wb = load_workbook(str(file_path), data_only=True)
    title = file_path.stem
    tables = []

    for sheet in wb.worksheets:
        # 병합 셀 정보 수집 (어떤 셀이 병합 범위의 왼쪽 상단인지)
        merge_map = {}        # (row, col) → 왼쪽 상단 값
        merge_skip = set()    # 병합으로 인해 건너뛸 셀
        for merged_range in sheet.merged_cells.ranges:
            top_left_val = sheet.cell(merged_range.min_row, merged_range.min_col).value
            for r in range(merged_range.min_row, merged_range.max_row + 1):
                for c in range(merged_range.min_col, merged_range.max_col + 1):
                    if r == merged_range.min_row and c == merged_range.min_col:
                        merge_map[(r, c)] = top_left_val
                    else:
                        merge_skip.add((r, c))

        # 모든 행의 데이터 수집
        all_rows = []
        for row_idx, row in enumerate(sheet.iter_rows(min_row=1), start=1):
            cell_values = []
            for col_idx, cell in enumerate(row, start=1):
                if (row_idx, col_idx) in merge_skip:
                    continue  # 병합된 셀의 하위 영역은 건너뛰기
                val = cell.value
                if val is None:
                    continue
                # 날짜 포맷팅
                if isinstance(val, datetime):
                    val = val.strftime('%Y-%m-%d')
                cell_values.append(str(val).strip())

            # 빈 행 건너뛰기
            non_empty = [v for v in cell_values if v]
            if non_empty:
                all_rows.append(non_empty)

        if not all_rows:
            continue

        # 제목 감지 (첫 1~2행에서 단일 셀이면 제목)
        start_idx = 0
        if len(all_rows[0]) == 1 and len(all_rows[0][0]) < 50:
            title = all_rows[0][0]
            start_idx = 1

        # 행 데이터를 HTML 테이블 구조로 변환
        rows_data = []
        for row_cells in all_rows[start_idx:]:
            if len(row_cells) == 1:
                rows_data.append({"label": None, "value": row_cells[0], "type": "text"})
            elif len(row_cells) == 2:
                row_type = "textarea" if '\n' in row_cells[1] or len(row_cells[1]) > 80 else "input"
                rows_data.append({"label": row_cells[0], "value": row_cells[1], "type": row_type})
            elif len(row_cells) >= 3:
                # 다열 데이터 → raw_cells로 전달하여 HTML에서 다열 테이블로 렌더링
                rows_data.append({
                    "label": None,
                    "value": None,
                    "type": "multi_cell",
                    "cells": row_cells
                })

        if rows_data:
            sheet_caption = sheet.title if len(wb.worksheets) > 1 else ""
            tables.append({"caption": sheet_caption, "rows": rows_data})

    wb.close()
    return title, tables, []


# ═════════════════════════════════════════════
# [4] PDF 파서
# ═════════════════════════════════════════════

def _parse_pdf(file_path: Path) -> tuple:
    """
    PDF 파일에서 텍스트와 표를 추출합니다.

    pdfplumber 라이브러리를 사용하여:
      - 페이지별 표(table) 자동 감지 및 추출
      - 표 외부의 일반 텍스트 추출
      - 이미지 기반 PDF는 지원 불가 (OCR 필요)

    Returns:
        (title, tables, paragraphs) 튜플
    """
    import pdfplumber
    import re as _re

    pdf = pdfplumber.open(str(file_path))
    title = file_path.stem
    tables = []
    extra_paragraphs = []
    title_found = False

    # ── CIDFont 감지: 첫 페이지 텍스트에 (cid:숫자) 패턴이 많으면 경고 ──
    first_text = (pdf.pages[0].extract_text() or "") if pdf.pages else ""
    cid_count = len(_re.findall(r'\(cid:\d+\)', first_text))
    if cid_count > 5:
        pdf.close()
        logger.warning("PDF에 CIDFont(커스텀 폰트)가 사용되어 텍스트 추출이 불가합니다.")
        raise ValueError(
            "이 PDF는 커스텀 폰트(CIDFont)를 사용하여 텍스트를 추출할 수 없습니다.\n"
            "해결 방법: 원본 HWP/DOCX 파일을 사용하거나, PDF를 '인쇄 > PDF로 저장'으로 다시 생성해 주세요."
        )

    for page_idx, page in enumerate(pdf.pages):
        # ── 표 추출 시도 ──
        page_tables = page.extract_tables() or []

        for table in page_tables:
            rows_data = []
            for row in table:
                # None 값을 빈 문자열로 변환
                cells = [str(c).strip() if c else '' for c in row]
                non_empty = [c for c in cells if c]

                if not non_empty:
                    continue

                if len(non_empty) == 1:
                    # 단일 셀 → 헤더 또는 텍스트
                    val = non_empty[0]
                    if not title_found and page_idx == 0 and len(val) < 30:
                        title = val
                        title_found = True
                        continue
                    rows_data.append({"label": None, "value": val, "type": "text"})
                elif len(non_empty) == 2:
                    label, value = non_empty[0], non_empty[1]
                    row_type = "textarea" if '\n' in value or len(value) > 80 else "input"
                    rows_data.append({"label": label, "value": value, "type": row_type})
                else:
                    label = non_empty[0]
                    value = ' | '.join(non_empty[1:])
                    rows_data.append({"label": label, "value": value, "type": "input"})

            if rows_data:
                tables.append({"caption": "", "rows": rows_data})

        # ── 표 밖 텍스트 추출 ──
        # 표 영역을 제외한 텍스트 가져오기
        text = page.extract_text() or ""
        if text.strip() and not page_tables:
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            # 첫 페이지 첫 줄을 제목으로
            if not title_found and page_idx == 0 and lines:
                title = lines[0]
                title_found = True
                lines = lines[1:]
            extra_paragraphs.extend(lines)

    pdf.close()

    # 표가 없고 텍스트만 있는 경우 → 라벨-값 감지 시도
    if not tables and extra_paragraphs:
        rows = _detect_label_value_pairs(extra_paragraphs)
        tables = [{"caption": "", "rows": rows}]
        extra_paragraphs = []

    return title, tables, extra_paragraphs


# ═════════════════════════════════════════════
# 통합 변환 메인 함수
# ═════════════════════════════════════════════

def convert(input_path: str, output_path: str = None) -> str:
    """
    파일 형식을 자동 감지하고 다우오피스 전자결재용 HTML로 변환합니다.

    이 함수 하나로 HWP, DOCX, Excel, PDF 모두 처리할 수 있습니다.

    Args:
        input_path: 변환할 원본 파일 경로
        output_path: 저장할 HTML 파일 경로 (생략 시 자동 생성)

    Returns:
        생성된 HTML 파일 경로 문자열
    """
    file_path = Path(input_path)
    ext = file_path.suffix.lower()

    # ── 파일 존재 여부 확인 ──
    if not file_path.exists():
        raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"지원하지 않는 파일 형식입니다: {ext}\n"
            f"지원 형식: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    # ── 출력 경로 자동 설정 ──
    if not output_path:
        output_path = file_path.with_suffix('.html')
    output_path = Path(output_path)

    logger.info("=" * 55)
    logger.info("다우오피스 전자결재 양식 변환 시작")
    logger.info("  입력 파일 : %s", file_path)
    logger.info("  파일 형식 : %s", ext.upper())
    logger.info("=" * 55)

    # ── DOCX: mammoth 라이브러리로 고품질 HTML 직접 생성 ──
    if ext == '.docx':
        logger.info("[1/2] mammoth로 DOCX → HTML 변환 중...")
        import mammoth as _mammoth
        from docx import Document as _DocxDoc

        # mammoth으로 HTML 변환 (테이블 구조/이미지 완벽 보존)
        with open(str(file_path), 'rb') as f:
            mammoth_result = _mammoth.convert_to_html(f)
        body_html = mammoth_result.value

        # 제목 감지 (python-docx 사용)
        _doc = _DocxDoc(str(file_path))
        title = ""
        for para in _doc.paragraphs:
            if para.style.name.startswith('Heading') or (
                para.alignment is not None and para.alignment == 1
            ):
                if para.text.strip():
                    title = para.text.strip()
                    break
        if not title:
            for para in _doc.paragraphs:
                if para.text.strip():
                    title = para.text.strip()
                    break
        if not title:
            stem = file_path.stem
            title = stem if not _re_common.match(r'^[0-9a-f]{8}$', stem) else "문서"

        # 다우오피스 양식 래퍼 적용
        # mammoth HTML의 <table>에 테두리/스타일 추가
        styled_body = body_html.replace(
            '<table>',
            '<table border="1" cellpadding="6" cellspacing="0" '
            'style="width: 100%; border-collapse: collapse; margin-bottom: 16px;">'
        ).replace(
            '<td>',
            '<td style="padding: 6px 10px;">'
        ).replace(
            '<td colspan',
            '<td style="padding: 6px 10px;" colspan'
        ).replace(
            '<td rowspan',
            '<td style="padding: 6px 10px;" rowspan'
        )

        html = (
            '<!-- 다우오피스 전자결재 양식 - mammoth + unified_converter.py 자동 생성 -->\n'
            '<div data-id="appContent">\n'
            '\n'
            '<div style="font-family: \'Malgun Gothic\', dotum, Arial, sans-serif; '
            'font-size: 10pt; line-height: 1.6; margin: 0 auto; max-width: 800px;">\n'
            '\n'
            '  <div data-id="appTitle" style="text-align: center; margin-bottom: 20px;">\n'
            '    <h2 style="font-size: 16pt; font-weight: bold; letter-spacing: 4px; '
            'border-bottom: 2px solid #333; padding-bottom: 8px;">\n'
            f'      {html_escape(title)}\n'
            '    </h2>\n'
            '  </div>\n'
            '\n'
            f'{styled_body}\n'
            '\n'
            '</div><!-- /font-family div -->\n'
            '</div><!-- /data-id="appContent" -->'
        )

        logger.info("[2/2] HTML 파일 저장 중...")
        output_path.write_text(html, encoding='utf-8')
        logger.info("       → 저장 완료: %s", output_path)
        logger.info("       → 크기: %s bytes", f"{len(html.encode('utf-8')):,}")
        logger.info("=" * 55)
        return str(output_path)

    # ── HWP/Excel/PDF: 기존 파서 호출 ──
    parser_map = {
        '.hwp':  _parse_hwp,
        '.xlsx': _parse_excel,
        '.xls':  _parse_excel,
        '.pdf':  _parse_pdf,
    }

    parser = parser_map[ext]
    logger.info("[1/3] %s 파일 파싱 중...", ext.upper())
    title, tables, paragraphs = parser(file_path)

    # 총 필드 수 계산
    total_fields = sum(len(t.get("rows", [])) for t in tables)
    logger.info("       → 제목: %s", title)
    logger.info("       → 테이블: %d개, 필드: %d개", len(tables), total_fields)

    # ── HTML 생성 ──
    logger.info("[2/3] 다우오피스 HTML 양식 생성 중...")
    html = build_daou_html(title, tables, paragraphs)

    # ── 파일 저장 ──
    logger.info("[3/3] HTML 파일 저장 중...")
    output_path.write_text(html, encoding='utf-8')
    logger.info("       → 저장 완료: %s", output_path)
    logger.info("       → 크기: %s bytes", f"{len(html.encode('utf-8')):,}")
    logger.info("=" * 55)

    return str(output_path)


# ─────────────────────────────────────────────
# 명령줄 실행
# ─────────────────────────────────────────────
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("사용법: python unified_converter.py <파일경로>")
        print()
        print("지원 형식:")
        print("  .hwp   - 한글(HWP) 문서")
        print("  .docx  - Microsoft Word 문서")
        print("  .xlsx  - Microsoft Excel 문서")
        print("  .pdf   - PDF 문서")
        print()
        print("예시:")
        print('  python unified_converter.py "(양식)차량사고경위서.hwp"')
        print('  python unified_converter.py "보고서.docx"')
        sys.exit(0)

    input_file = sys.argv[1]
    out_file = sys.argv[2] if len(sys.argv) > 2 else None

    try:
        result = convert(input_file, out_file)
        print(f"\n변환 완료! 결과 파일: {result}")
    except Exception as e:
        logger.error("변환 실패: %s", e)
        sys.exit(1)
