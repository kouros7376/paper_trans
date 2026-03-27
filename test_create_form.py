# -*- coding: utf-8 -*-
"""
테스트용 가상 전자결재양식(휴가신청서) DOCX 파일 생성 스크립트
- 실제 회사에서 쓰이는 휴가신청서 양식을 모사합니다.
- 이 파일을 unified_converter로 변환하여 원본과 비교합니다.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path


def set_cell_shading(cell, color_hex: str):
    """셀 배경색을 설정하는 헬퍼 함수"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def set_cell_text(cell, text: str, bold=False, size=10, align='center'):
    """셀 내 텍스트를 설정하는 헬퍼 함수"""
    cell.text = ""
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '맑은 고딕'
    run.bold = bold


def create_leave_request_form(output_path: str):
    """가상 휴가신청서 DOCX 파일을 생성합니다."""
    doc = Document()

    # ── 문서 제목 ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("휴 가 신 청 서")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = '맑은 고딕'

    # ── 1. 신청자 정보 테이블 ──
    header1 = doc.add_paragraph()
    run1 = header1.add_run("1. 신청자 정보")
    run1.font.size = Pt(11)
    run1.font.bold = True

    # 4열 테이블: 라벨 | 값 | 라벨 | 값
    table1 = doc.add_table(rows=3, cols=4)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 소속 행
    set_cell_text(table1.cell(0, 0), "소 속", bold=True)
    set_cell_shading(table1.cell(0, 0), "E6E8EB")
    set_cell_text(table1.cell(0, 1), "전략설계1본부", align='left')
    set_cell_text(table1.cell(0, 2), "직 위", bold=True)
    set_cell_shading(table1.cell(0, 2), "E6E8EB")
    set_cell_text(table1.cell(0, 3), "대리", align='left')

    # 성명 행
    set_cell_text(table1.cell(1, 0), "성 명", bold=True)
    set_cell_shading(table1.cell(1, 0), "E6E8EB")
    set_cell_text(table1.cell(1, 1), "홍길동", align='left')
    set_cell_text(table1.cell(1, 2), "사 번", bold=True)
    set_cell_shading(table1.cell(1, 2), "E6E8EB")
    set_cell_text(table1.cell(1, 3), "220125", align='left')

    # 연락처 행
    set_cell_text(table1.cell(2, 0), "연락처", bold=True)
    set_cell_shading(table1.cell(2, 0), "E6E8EB")
    set_cell_text(table1.cell(2, 1), "010-1234-5678", align='left')
    set_cell_text(table1.cell(2, 2), "이메일", bold=True)
    set_cell_shading(table1.cell(2, 2), "E6E8EB")
    set_cell_text(table1.cell(2, 3), "hong@haema.co.kr", align='left')

    doc.add_paragraph()  # 빈 줄

    # ── 2. 휴가 내용 테이블 ──
    header2 = doc.add_paragraph()
    run2 = header2.add_run("2. 휴가 내용")
    run2.font.size = Pt(11)
    run2.font.bold = True

    table2 = doc.add_table(rows=4, cols=2)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 휴가 종류
    set_cell_text(table2.cell(0, 0), "휴가 종류", bold=True)
    set_cell_shading(table2.cell(0, 0), "E6E8EB")
    set_cell_text(table2.cell(0, 1), "연차 휴가", align='left')

    # 휴가 기간
    set_cell_text(table2.cell(1, 0), "휴가 기간", bold=True)
    set_cell_shading(table2.cell(1, 0), "E6E8EB")
    set_cell_text(table2.cell(1, 1), "2026년 04월 01일 ~ 2026년 04월 03일 (3일간)", align='left')

    # 휴가 사유
    set_cell_text(table2.cell(2, 0), "휴가 사유", bold=True)
    set_cell_shading(table2.cell(2, 0), "E6E8EB")
    set_cell_text(table2.cell(2, 1), "가족 여행 (제주도)", align='left')

    # 비상 연락처
    set_cell_text(table2.cell(3, 0), "비상 연락처", bold=True)
    set_cell_shading(table2.cell(3, 0), "E6E8EB")
    set_cell_text(table2.cell(3, 1), "010-1234-5678 (본인)", align='left')

    doc.add_paragraph()  # 빈 줄

    # ── 3. 업무 인계 사항 테이블 ──
    header3 = doc.add_paragraph()
    run3 = header3.add_run("3. 업무 인계 사항")
    run3.font.size = Pt(11)
    run3.font.bold = True

    table3 = doc.add_table(rows=2, cols=2)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_cell_text(table3.cell(0, 0), "인계 대상자", bold=True)
    set_cell_shading(table3.cell(0, 0), "E6E8EB")
    set_cell_text(table3.cell(0, 1), "김철수 과장 (정보시스템팀)", align='left')

    set_cell_text(table3.cell(1, 0), "인계 내용", bold=True)
    set_cell_shading(table3.cell(1, 0), "E6E8EB")
    set_cell_text(table3.cell(1, 1),
                  "1. 서버 모니터링 점검 (매일 오전 9시)\n"
                  "2. 다우오피스 API 연동 테스트 진행\n"
                  "3. 방화벽 정책 변경 요청 건 처리",
                  align='left')

    doc.add_paragraph()  # 빈 줄

    # ── 제출 문구 ──
    submit_text = doc.add_paragraph()
    submit_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = submit_text.add_run("위와 같이 휴가를 신청합니다.")
    run_s.font.size = Pt(10)

    date_text = doc.add_paragraph()
    date_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_d = date_text.add_run("2026년  3월  27일")
    run_d.font.size = Pt(10)

    # ── 서명란 ──
    sign_text = doc.add_paragraph()
    sign_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sign = sign_text.add_run("신청자 :  홍 길 동  (인)")
    run_sign.font.size = Pt(10)

    # 파일 저장
    doc.save(output_path)
    print(f"가상 전자결재양식 생성 완료: {output_path}")


if __name__ == "__main__":
    output = str(Path(__file__).parent / "(테스트)휴가신청서.docx")
    create_leave_request_form(output)
